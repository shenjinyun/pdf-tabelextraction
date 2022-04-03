import itertools
from pdfplumber import utils
from operator import itemgetter
from docx import Document

DEFAULT_EXTRACT_X_TOLERANCE = 3
DEFAULT_EXTRACT_Y_TOLERANCE = 3

class CellGroup(object):
    def __init__(self, cells):
        self.cells = cells
        self.bbox = (
            min(map(itemgetter(0), filter(None, cells))),
            min(map(itemgetter(1), filter(None, cells))),
            max(map(itemgetter(2), filter(None, cells))),
            max(map(itemgetter(3), filter(None, cells))),
        )

class Row(CellGroup):
    pass

class Table(object):
    def __init__(self, page, cells):
        self.page = page
        self.cells = cells
        self.bbox = (
            min(map(itemgetter(0), cells)),
            min(map(itemgetter(1), cells)),
            max(map(itemgetter(2), cells)),
            max(map(itemgetter(3), cells)),
        )

    @property
    def rows(self):
        _sorted = sorted(self.cells, key=itemgetter(1, 0)) 
        xs = list(sorted(set(map(itemgetter(0), self.cells))))
        rows = []
        for y, row_cells in itertools.groupby(_sorted, itemgetter(1)):
            xdict = dict((cell[0], cell) for cell in row_cells)
            row = Row([ xdict.get(x) for x in xs ])
            rows.append(row)
        return rows

    def extract(self,
        extract_x_tolerance=DEFAULT_EXTRACT_X_TOLERANCE,
        extract_y_tolerance=DEFAULT_EXTRACT_Y_TOLERANCE):

        chars = self.page.chars
        table_arr = []
        def char_in_bbox(char, bbox):
            v_mid = (char["top"] + char["bottom"]) / 2
            h_mid = (char["x0"] + char["x1"]) / 2
            x0, top, x1, bottom = bbox
            return (
                (h_mid >= x0) and
                (h_mid < x1) and
                (v_mid >= top) and
                (v_mid < bottom)
            )

        for row in self.rows:
            arr = []
            row_chars = [ char for char in chars
                if char_in_bbox(char, row.bbox) ]

            for cell in row.cells:
                if cell == None:
                    cell_text = None
                else:
                    cell_chars = [ char for char in row_chars
                        if char_in_bbox(char, cell) ]

                    if len(cell_chars):
                        cell_text = utils.extract_text(cell_chars,
                            x_tolerance=extract_x_tolerance,
                            y_tolerance=extract_y_tolerance).strip()
                    else:
                        cell_text = ""
                arr.append(cell_text)
            table_arr.append(arr)

        return table_arr
   
def cells_to_tables(page, cells):
    def get_table_cells(cells):
        '''
        将所有的cell分组，属于一个table的cell分到同一组
        '''

        def bbox_to_corners(bbox): # 根据bbox得到cell的四个顶点坐标
                x0, top, x1, bottom = bbox
                return list(itertools.product((x0, x1), (top, bottom)))
                        

        cells = [ {
        
                "available": True,
                "bbox": bbox,
                "corners": bbox_to_corners(bbox)
                } for bbox in cells ]

        def init_new_table():
                return { "corners": set([]), "cells": [] }

        def assign_cell(cell, table):
                table["corners"] = table["corners"].union(set(cell["corners"]))
                table["cells"].append(cell["bbox"])
                cell["available"] = False

        n_cells = len(cells)
        n_assigned = 0
        tables = []
        current_table = init_new_table()

        while True:
                initial_cell_count = len(current_table["cells"])
                for i, cell in enumerate(filter(itemgetter("available"), cells)):
                        if len(current_table["cells"]) == 0:
                                assign_cell(cell, current_table)
                                n_assigned += 1
                        else:
                                corner_count = sum(c in current_table["corners"]
                                        for c in cell["corners"])
                                if corner_count > 0 and cell["available"]:
                                        assign_cell(cell, current_table)
                                        n_assigned += 1
                if n_assigned == n_cells:
                        break
                if len(current_table["cells"]) == initial_cell_count:
                        tables.append(current_table)
                        current_table = init_new_table()

        if len(current_table["cells"]):
                tables.append(current_table)

        _sorted = sorted(tables, key = lambda t: min(t["corners"]))
        filtered = [ t["cells"] for t in _sorted if len(t["cells"]) > 1 ]
        return filtered

    table_cells = get_table_cells(cells)

    tables = [ Table(page, t)
               for t in table_cells ]

    return [ table.extract() for table in tables ]

def show_tables(tables):
    table_num = len(tables)

    for i in range(0, table_num):
        print("\ntable_" + str(i + 1) + ":")
        row_num = len(tables[i])
        for j in range(0, row_num):
            print(tables[i][j])
	
def save_tables(tables, file_name):
    document = Document()
    table_num = len(tables)

    for i in range(0, table_num):
        document.add_paragraph("table_" + str(i + 1) + ":")

        current_table = tables[i]
        row_num = len(current_table)
        column_num = len(current_table[0])

        docx_table = document.add_table(rows = row_num, cols = column_num, style = "Table Grid")
        for j in range(0, row_num):
            for k in range(0, column_num):
                if isinstance(current_table[j][k], str):
                    docx_table.cell(j, k).text = current_table[j][k]
                else:
                    docx_table.cell(j, k).text = " "

        document.add_paragraph("")

        document.save(file_name.strip(".pdf") + "_tables.docx")
    print("保存成功！")
			
	
